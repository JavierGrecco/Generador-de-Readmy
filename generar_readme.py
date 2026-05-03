#!/usr/bin/env python3
"""
Generador de README.md para GitHub.

Convierte un archivo .txt o .docx en un README.md profesional,
añadiendo badges, tabla de contenidos y emojis.
Incluye logging a archivo, validación de entrada y modo debug del AST.

Decisiones de diseño
--------------------
* ¿Por qué un AST y no regex?
  Procesar Markdown con expresiones regulares es frágil. Los bloques de código
  anidan, los fences pueden estar balanceados o no, y las tablas mezclan pipes
  con pipe de Markdown. Un AST (aunque simplificado) permite razonar sobre la
  estructura del documento y aplicar transformaciones sin romper otros bloques.

* ¿Por qué este orden en el pipeline?
  1) Limpieza – eliminar ruido que pueda confundir al parser.
  2) Insertar encabezados – normalizar títulos que no empiezan con ##.
  3) Normalizar comandos – envolver cmd/bash/text en fences.
  4) Detectar bloques – agrupar líneas en bloques tipificados.
  5) Construir AST – armar el árbol sintáctico.
  6) Visitores – enriquecer con emojis y normalizar duplicados.
  7) Renderizar – generar Markdown limpio.
  Este orden minimiza las interferencias entre etapas. Por ejemplo, si
  envolvemos comandos después de detectar bloques, rompemos la detección de
  fences. Si insertamos encabezados después de normalizar comandos, podemos
  romper bloques de código.

* ¿Por qué forzar algunos encabezados?
  Documentos que vienen de Word o texto plano suelen tener secciones como
  "Windows", "Linux", "Uso" sin el prefijo ##. El script las convierte
  automáticamente para garantizar una TOC completa.

Edge cases resueltos a mano
---------------------------
* Las líneas de licencia (MIT © ...) se filtraban para que no se convirtieran
  en encabezados. Al principio solo se miraba el prefijo MIT, pero luego vi que
  también aparecían con otros formatos, así que agregué la detección de ©.
* Las líneas con pipes (|) son tablas, pero también pueden ser texto normal en
  bloques de código. Para no romperlas, el filtro solo actúa fuera de fences.
* Los bloques de código con comentarios # dentro de bash/cmd se cortaban antes
  de tiempo. Tuve que agregar una lista de lenguajes de shell donde el #
  es parte del lenguaje y no un encabezado.
* El árbol de archivos se fragmentaba porque había líneas de comentario (#)
  entre las ramas. Ahora el detector de árbol las incluye siempre que sean
  cortas y no empiecen con ##.
* La sección "Instalación" desaparecía porque era un nodo sin hijos y el
  normalizador la eliminaba. Agregué una lista de secciones agrupadoras que
  nunca se borran aunque estén vacías.

Autor: Javier Grecco (https://github.com/JavierGrecco)
Licencia: MIT
"""

import argparse
import logging
import os
import re
import sys
from pathlib import Path
from typing import Dict, List, Optional, Tuple

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
def configurar_logging() -> None:
    logger = logging.getLogger()
    logger.setLevel(logging.DEBUG)

    archivo = logging.FileHandler("generar_readme.log", encoding="utf-8")
    archivo.setLevel(logging.DEBUG)
    archivo.setFormatter(logging.Formatter("[%(levelname)s] %(asctime)s - %(message)s"))

    consola = logging.StreamHandler()
    consola.setLevel(logging.INFO)
    consola.setFormatter(logging.Formatter("%(levelname)s: %(message)s"))

    logger.addHandler(archivo)
    logger.addHandler(consola)


# ---------------------------------------------------------------------------
# Configuración
# ---------------------------------------------------------------------------
BADGES_TECNOLOGIAS = {
    "python": ("Python", "3776AB", "python"),
    "javascript": ("JavaScript", "F7DF1E", "javascript"),
    "typescript": ("TypeScript", "3178C6", "typescript"),
    "react": ("React", "61DAFB", "react"),
    "node": ("Node.js", "339933", "node.js"),
    "docker": ("Docker", "2496ED", "docker"),
    "aws": ("AWS", "FF9900", "amazonwebservices"),
    "gcp": ("Google Cloud", "4285F4", "googlecloud"),
    "shell": ("Shell", "4EAA25", "gnubash"),
    "html": ("HTML5", "E34F26", "html5"),
    "css": ("CSS3", "1572B6", "css3"),
    "java": ("Java", "007396", "java"),
    "rust": ("Rust", "000000", "rust"),
    "go": ("Go", "00ADD8", "go"),
    "ruby": ("Ruby", "CC342D", "ruby"),
    "php": ("PHP", "777BB4", "php"),
    "c++": ("C++", "00599C", "cplusplus"),
    ".net": (".NET", "512BD4", "dotnet"),
    "windows": ("Windows", "0078D6", "windows"),
    "linux": ("Linux", "FCC624", "linux"),
    "macos": ("macOS", "000000", "apple"),
    "excel": ("Excel", "217346", "microsoftexcel"),
    "csv": ("CSV", "217346", "csv"),
    "pdf": ("PDF", "FF0000", "adobeacrobatreader"),
}

LICENCIAS_CONOCIDAS = {
    "mit": ("MIT", "yellow"),
    "apache": ("Apache 2.0", "blue"),
    "gpl": ("GPL v3", "blue"),
    "lgpl": ("LGPL v3", "blue"),
    "bsd": ("BSD", "orange"),
    "mpl": ("MPL 2.0", "orange"),
    "unlicense": ("Unlicense", "blue"),
}

EMOJI_SECCIONES = {
    "introducción": ":book:",
    "características": ":sparkles:",
    "instalación por sistema operativo": ":computer:",
    "instalación": ":wrench:",
    "requisitos": ":clipboard:",
    "estructura de archivos": ":file_folder:",
    "uso": ":rocket:",
    "resultados": ":package:",
    "personalización": ":art:",
    "formato de los datos": ":bar_chart:",
    "archivo de descripciones": ":page_facing_up:",
    "notas adicionales": ":memo:",
    "contribución": ":handshake:",
    "licencia": ":scroll:",
    "créditos": ":star:",
    "windows": ":computer:",
    "linux": ":computer:",
    "macos": ":computer:",
    "debian": ":computer:",
    "ubuntu": ":computer:",
    "debian/ubuntu": ":computer:",
    "logging": ":memo:",
    "tests": ":test_tube:",
}

FORCED_HEADERS = {
    "instalación por sistema operativo", "windows", "linux", "macos",
    "estructura de archivos", "formato de los datos",
    "archivo de descripciones por categoría (opcional)",
    "uso", "resultados", "personalización", "notas adicionales",
    "cómo funciona", "como funciona", "características", "caracteristicas",
    "requisitos", "instalación", "instalacion", "logging", "tests", "licencia",
    "debian/ubuntu", "debian", "ubuntu",
}

INDICADORES_CODIGO = {
    "cmd", "bash", "sh", "powershell", "python", "shell",
    "text", "yml", "yaml", "json",
}

PALABRAS_PROHIBIDAS = {
    "cmd", "bash", "text", "sh", "powershell", "python", "shell",
    "pip", "python3", "sudo", "apt", "brew", "npm", "git", "cd",
    "--txt", "--docx", "--auto", "--license", "--logo",
    "--no-toc", "--no-credits", "--no-mandatory", "--watch",
    "-o", "--output", "pytest",
}

COMANDOS_CONOCIDOS = {
    "pip", "python", "python3", "cd", "sudo", "npm", "git", "brew", "apt", "pytest",
}

SECCIONES_AGRUPADORAS = {"instalación", "instalacion"}


# ---------------------------------------------------------------------------
# Nodo del AST
# ---------------------------------------------------------------------------
class Nodo:
    """Representa un elemento del árbol sintáctico del documento."""

    def __init__(self, tipo: str, contenido: str = "", meta: Optional[Dict] = None):
        self.tipo = tipo
        self.contenido = contenido
        self.hijos: List["Nodo"] = []
        self.meta = meta or {}

    def agregar_hijo(self, nodo: "Nodo") -> None:
        self.hijos.append(nodo)


# ---------------------------------------------------------------------------
# Funciones auxiliares de clasificación de líneas
# ---------------------------------------------------------------------------
def _linea_es_tabla(linea: str) -> bool:
    return "|" in linea

def _linea_es_arbol(linea: str) -> bool:
    return any(c in linea for c in "│├└─")

def _linea_es_licencia(linea: str) -> bool:
    if re.match(r"^(MIT|Apache|GPL|BSD|MPL|ISC)\b", linea, re.IGNORECASE):
        return True
    return "©" in linea and "github.com" in linea.lower()

def _linea_prohibida(linea: str) -> bool:
    primera = linea.split()[0].lower().lstrip("-")
    return primera in PALABRAS_PROHIBIDAS or linea.lower() in PALABRAS_PROHIBIDAS


# ---------------------------------------------------------------------------
# Validación de entrada
# ---------------------------------------------------------------------------
def validar_archivo(ruta: str) -> Optional[str]:
    """Verifica que el archivo exista, sea legible y tenga contenido."""
    camino = Path(ruta)
    if not camino.exists():
        logging.error("El archivo '%s' no existe.", ruta)
        return None
    if not camino.is_file():
        logging.error("'%s' no es un archivo.", ruta)
        return None
    if camino.stat().st_size == 0:
        logging.error("El archivo '%s' está vacío.", ruta)
        return None

    # Los archivos .docx son binarios, omitimos comprobación de texto
    if camino.suffix.lower() == ".docx":
        return str(camino)

    try:
        with open(camino, "r", encoding="utf-8") as f:
            f.read(1)
    except (OSError, UnicodeDecodeError) as e:
        logging.error("No se pudo leer '%s': %s", ruta, e)
        return None
    return str(camino)


# ---------------------------------------------------------------------------
# Lectura de archivos
# ---------------------------------------------------------------------------
def leer_docx(ruta: str) -> Optional[str]:
    """Extrae el texto y las tablas de un documento .docx."""
    if DocxDocument is None:
        logging.error("python-docx no está instalado. Instalalo con: pip install python-docx")
        return None
    try:
        doc = DocxDocument(ruta)
        partes = []
        for p in doc.paragraphs:
            if p.text.strip():
                partes.append(p.text.strip())
        for table in doc.tables:
            filas = [[celda.text.strip().replace("\n", " ") for celda in row.cells] for row in table.rows]
            if filas:
                max_cols = max(len(f) for f in filas)
                filas = [f + [""] * (max_cols - len(f)) for f in filas]
                lineas = [
                    "| " + " | ".join(filas[0]) + " |",
                    "| " + " | ".join(["---"] * max_cols) + " |",
                ]
                for f in filas[1:]:
                    lineas.append("| " + " | ".join(f) + " |")
                partes.append("\n".join(lineas))
        resultado = "\n\n".join(partes)
        if not resultado.strip():
            partes = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            resultado = "\n".join(partes)
        logging.info("Documento DOCX leído correctamente.")
        return resultado or None
    except Exception as error:
        logging.error("Error al leer DOCX %s: %s", ruta, error)
        return None


def leer_txt(ruta: str) -> Optional[str]:
    """Lee un archivo de texto plano."""
    try:
        contenido = Path(ruta).read_text(encoding="utf-8")
        logging.info("Archivo TXT leído correctamente.")
        return contenido
    except Exception as error:
        logging.error("Error al leer TXT %s: %s", ruta, error)
        return None


def buscar_archivo() -> Optional[Path]:
    """Busca en el directorio actual un archivo .txt o .docx adecuado."""
    txts = list(Path(".").glob("*.txt"))
    docxs = list(Path(".").glob("*.docx")) if DocxDocument else []
    candidatos = [p for p in txts + docxs if not p.name.startswith("~$")]
    if not candidatos:
        logging.warning("No se encontraron archivos .txt o .docx en el directorio.")
        return None
    for ruta in candidatos:
        nombre = ruta.name.lower()
        if any(p in nombre for p in ["readme", "descripcion", "descripción", "catalogo"]):
            logging.info("Archivo detectado por nombre: %s", ruta)
            return ruta
    mas_reciente = max(candidatos, key=lambda p: p.stat().st_mtime)
    logging.info("Archivo más reciente detectado: %s", mas_reciente)
    return mas_reciente


# ---------------------------------------------------------------------------
# Limpieza
# ---------------------------------------------------------------------------
def sanitizar(texto: str) -> str:
    """Elimina badges existentes, líneas espurias y patrones de subrayado de Word."""
    texto = re.sub(r'!\[\]\(https://img\.shields\.io[^\)]*\)', '', texto)
    lineas = []
    for linea in texto.splitlines():
        limpia = linea.strip()
        if limpia in ("Copiar", "Descargar"):
            continue
        limpia = re.sub(r"\[(.*?)\]\{\.underline\}", r"\1", limpia)
        lineas.append(limpia)
    return "\n".join(lineas)


# ---------------------------------------------------------------------------
# Detección de encabezados
# ---------------------------------------------------------------------------
def es_header_semantico(linea: str) -> bool:
    """Indica si una línea parece un título que debería convertirse en encabezado."""
    despojada = linea.strip()
    if not despojada or despojada.startswith("#"):
        return False
    if _linea_es_tabla(despojada) or _linea_es_arbol(despojada) or _linea_es_licencia(despojada):
        return False
    if _linea_prohibida(despojada):
        return False
    if re.match(r'^:\w+:\s+\w+', despojada):
        return True
    if 1 <= len(despojada.split()) <= 6 and not despojada.endswith('.'):
        if despojada[0].islower() and not any(c.isupper() for c in despojada):
            return False
        return True
    return False


def forzar_titulos(texto: str) -> str:
    """Convierte líneas con aspecto de título en encabezados ##."""
    lineas = texto.splitlines()
    resultado = []
    dentro_de_codigo = False

    for i, linea in enumerate(lineas):
        despojada = linea.strip()

        if despojada.startswith("```"):
            dentro_de_codigo = not dentro_de_codigo
            resultado.append(linea)
            continue
        if dentro_de_codigo:
            resultado.append(linea)
            continue
        if despojada.startswith("#"):
            resultado.append(linea)
            continue
        if not despojada:
            resultado.append(linea)
            continue
        if _linea_es_tabla(despojada) or _linea_es_arbol(despojada) or _linea_es_licencia(despojada):
            resultado.append(linea)
            continue
        if despojada.lower() in FORCED_HEADERS:
            resultado.append(f"## {despojada}")
            continue
        if es_header_semantico(despojada):
            prev_vacia = (i == 0) or (not lineas[i - 1].strip())
            if prev_vacia:
                siguiente_es_codigo = False
                if i + 1 < len(lineas):
                    sig = lineas[i + 1].strip()
                    if sig and (
                        sig.lower() in INDICADORES_CODIGO
                        or sig.split()[0].lower() in PALABRAS_PROHIBIDAS
                    ):
                        siguiente_es_codigo = True
                if not siguiente_es_codigo:
                    resultado.append(f"## {despojada}")
                    continue
        resultado.append(linea)

    # Post-procesado: insertar "## Instalación" si Windows/Linux/macOS aparecen huérfanos
    final = []
    instalacion_insertada = False
    for i, linea in enumerate(resultado):
        despojada = linea.strip()
        if despojada.startswith("## ") and despojada[3:].strip().lower() in {
            "windows", "linux", "macos", "debian/ubuntu", "debian", "ubuntu"
        }:
            if not instalacion_insertada:
                encontrado = False
                for j in range(i - 1, -1, -1):
                    if resultado[j].strip().startswith("## "):
                        if "instalación" in resultado[j].strip().lower() or "instalacion" in resultado[j].strip().lower():
                            encontrado = True
                        break
                if not encontrado:
                    final.append("## Instalación")
                    instalacion_insertada = True
        final.append(linea)
    return "\n".join(final)


# ---------------------------------------------------------------------------
# Normalización de comandos
# ---------------------------------------------------------------------------
def envolver_comandos(texto: str) -> str:
    """Envuelve indicadores de código sueltos (cmd, bash, text...) en fences."""
    lineas = texto.splitlines()
    resultado = []
    i = 0
    while i < len(lineas):
        despojada = lineas[i].strip().lower()
        if despojada in INDICADORES_CODIGO and not despojada.startswith("```"):
            lenguaje = despojada
            i += 1
            codigo = []
            if lenguaje == "text":
                while i < len(lineas):
                    actual = lineas[i].strip()
                    if actual.startswith("##") or actual.lower() in FORCED_HEADERS:
                        break
                    if actual == "" and i + 1 < len(lineas) and lineas[i + 1].strip() == "":
                        break
                    codigo.append(lineas[i])
                    i += 1
            else:
                while i < len(lineas):
                    actual = lineas[i].strip()
                    if actual.startswith("##") or actual.lower() in FORCED_HEADERS or actual.lower() in INDICADORES_CODIGO:
                        break
                    if actual == "" and i + 1 < len(lineas):
                        sig = lineas[i + 1].strip()
                        if sig and not sig.startswith("#") and sig.lower() not in INDICADORES_CODIGO:
                            if any(sig.lower().startswith(c) for c in COMANDOS_CONOCIDOS) or re.match(r'^[\$>#]', sig):
                                codigo.append(lineas[i])
                                i += 1
                                continue
                            else:
                                codigo.append(lineas[i])
                                i += 1
                                break
                    codigo.append(lineas[i])
                    i += 1
            if codigo:
                while codigo and codigo[-1].strip() == "":
                    codigo.pop()
                if codigo and not any(codigo[-1].strip().startswith(c) for c in COMANDOS_CONOCIDOS):
                    ultima = codigo.pop()
                    resultado.append(f"```{lenguaje}")
                    resultado.extend(codigo)
                    resultado.append("```")
                    resultado.append(ultima)
                    continue
                resultado.append(f"```{lenguaje}")
                resultado.extend(codigo)
                resultado.append("```")
            continue
        resultado.append(lineas[i])
        i += 1
    return "\n".join(resultado)


def reparar_fences(md: str) -> str:
    """Cierra cualquier bloque de código que haya quedado abierto."""
    if md.count("```") % 2 != 0:
        md += "\n```"
    return md


# ---------------------------------------------------------------------------
# Detección de bloques y construcción del AST
# ---------------------------------------------------------------------------
def detectar_bloques(lineas: List[str]) -> List[Tuple[str, any]]:
    """Convierte la lista de líneas en bloques tipificados."""
    bloques = []
    i = 0
    while i < len(lineas):
        linea = lineas[i]
        despojada = linea.strip()

        if despojada.startswith("```"):
            lenguaje = despojada[3:].strip()
            i += 1
            codigo = []
            while i < len(lineas) and not lineas[i].strip().startswith("```"):
                codigo.append(lineas[i])
                i += 1
            if i < len(lineas) and lineas[i].strip().startswith("```"):
                i += 1
            bloques.append(("CODE_BLOCK", lenguaje, codigo))
            continue

        if "|" in despojada and not despojada.startswith("```"):
            tabla_lineas = [linea]
            i += 1
            while i < len(lineas) and "|" in lineas[i]:
                tabla_lineas.append(lineas[i])
                i += 1
            bloques.append(("TABLE_RAW", tabla_lineas))
            continue

        if any(c in despojada for c in "│├└─"):
            arbol_lineas = [linea]
            i += 1
            while i < len(lineas):
                actual = lineas[i].strip()
                if any(c in actual for c in "│├└─") or actual == "" or (
                    actual.startswith("#") and len(actual) < 80 and not actual.startswith("##")
                ):
                    if actual and not actual.startswith("##"):
                        arbol_lineas.append(lineas[i])
                    elif actual == "":
                        if arbol_lineas and arbol_lineas[-1].strip() != "":
                            arbol_lineas.append(lineas[i])
                    i += 1
                else:
                    break
            while arbol_lineas and arbol_lineas[-1].strip() == "":
                arbol_lineas.pop()
            bloques.append(("TREE", arbol_lineas))
            continue

        if despojada.startswith(("- ", "* ")):
            items = [despojada[2:].strip()]
            i += 1
            while i < len(lineas) and lineas[i].strip().startswith(("- ", "* ")):
                items.append(lineas[i].strip()[2:].strip())
                i += 1
            bloques.append(("LIST", items))
            continue

        bloques.append(("LINE", linea))
        i += 1

    return bloques


def construir_ast(bloques: List[Tuple[str, any]]) -> Nodo:
    """Arma el árbol sintáctico del documento a partir de los bloques."""
    raiz = Nodo("DOCUMENTO")
    pila = [raiz]
    buffer_parrafo: List[str] = []

    def vaciar_parrafo():
        if buffer_parrafo:
            pila[-1].agregar_hijo(Nodo("PARRAFO", contenido="\n".join(buffer_parrafo)))
            buffer_parrafo.clear()

    for bloque in bloques:
        tipo = bloque[0]
        if tipo == "CODE_BLOCK":
            vaciar_parrafo()
            _, lenguaje, codigo = bloque
            pila[-1].agregar_hijo(
                Nodo("BLOQUE_CODIGO", contenido="\n".join(codigo), meta={"lenguaje": lenguaje})
            )
        elif tipo == "TABLE_RAW":
            vaciar_parrafo()
            pila[-1].agregar_hijo(Nodo("TABLA_CRUDA", contenido="\n".join(bloque[1])))
        elif tipo == "LIST":
            vaciar_parrafo()
            nodo_lista = Nodo("LISTA")
            for item in bloque[1]:
                nodo_lista.agregar_hijo(Nodo("ITEM_LISTA", contenido=item))
            pila[-1].agregar_hijo(nodo_lista)
        elif tipo == "TREE":
            vaciar_parrafo()
            pila[-1].agregar_hijo(Nodo("ARBOL", contenido="\n".join(bloque[1])))
        elif tipo == "LINE":
            contenido = bloque[1].strip()
            if contenido == "":
                vaciar_parrafo()
            elif contenido.startswith("#"):
                vaciar_parrafo()
                nivel = contenido.count("#", 0, contenido.index(" ") if " " in contenido else len(contenido))
                seccion = Nodo("SECCION", contenido=bloque[1], meta={"nivel": nivel})
                while len(pila) > 1 and pila[-1].tipo == "SECCION" and pila[-1].meta.get("nivel", 1) >= nivel:
                    pila.pop()
                pila[-1].agregar_hijo(seccion)
                pila.append(seccion)
            else:
                buffer_parrafo.append(bloque[1])
    vaciar_parrafo()
    return raiz


# ---------------------------------------------------------------------------
# Visitores del AST
# ---------------------------------------------------------------------------
def walk(nodo: Nodo, funcion: callable) -> None:
    """Recorre el AST en profundidad y aplica una función a cada nodo."""
    funcion(nodo)
    for hijo in nodo.hijos:
        walk(hijo, funcion)


def visitor_emojis(nodo: Nodo) -> None:
    """Asigna emojis a las secciones del documento según su título."""
    if nodo.tipo != "SECCION":
        return
    match = re.match(r"^(#+)\s(.*)", nodo.contenido)
    if not match:
        return
    numerales, titulo = match.group(1), match.group(2)
    if re.match(r"^:\w+:", titulo.strip()):
        return
    for clave, emoji in EMOJI_SECCIONES.items():
        if clave in titulo.lower():
            nodo.contenido = f"{numerales} {emoji} {titulo}"
            return


# ---------------------------------------------------------------------------
# Normalización y validación del AST
# ---------------------------------------------------------------------------
def normalizar_ast(raiz: Nodo) -> Nodo:
    """Elimina secciones duplicadas y limpia nodos vacíos, excepto agrupadores como Instalación."""
    seen = set()
    nuevos = []
    for hijo in raiz.hijos:
        if hijo.tipo == "SECCION":
            titulo_low = hijo.contenido.lower()
            if "tabla de contenidos" in titulo_low:
                if "toc" in seen:
                    continue
                seen.add("toc")
            titulo_norm = re.sub(r":[a-z_]+:", "", titulo_low).strip()
            if titulo_norm in seen:
                continue
            seen.add(titulo_norm)
            if not hijo.hijos and len(hijo.contenido.split()) <= 4:
                if not any(p in titulo_low for p in SECCIONES_AGRUPADORAS):
                    continue
        nuevos.append(hijo)
    raiz.hijos = nuevos
    return raiz


def validar_estructura(raiz: Nodo) -> bool:
    """Comprueba que el AST tenga al menos una sección y no más de una TOC."""
    secciones = [c for c in raiz.hijos if c.tipo == "SECCION"]
    if not secciones:
        return False
    return sum(1 for s in secciones if "tabla de contenidos" in s.contenido.lower()) <= 1


# ---------------------------------------------------------------------------
# Renderizador Markdown
# ---------------------------------------------------------------------------
def renderizar_md(nodo: Nodo) -> str:
    """Convierte el AST en una cadena Markdown."""
    if nodo.tipo == "DOCUMENTO":
        return "\n".join(renderizar_md(h) for h in nodo.hijos).strip() + "\n"
    if nodo.tipo == "SECCION":
        cuerpo = "\n\n".join(renderizar_md(h) for h in nodo.hijos)
        return f"{nodo.contenido}\n\n{cuerpo}" if cuerpo else nodo.contenido
    if nodo.tipo == "PARRAFO":
        return nodo.contenido
    if nodo.tipo == "LISTA":
        return "\n".join(f"- {renderizar_md(h)}" for h in nodo.hijos)
    if nodo.tipo == "BLOQUE_CODIGO":
        lenguaje = nodo.meta.get("lenguaje", "")
        return f"```{lenguaje}\n{nodo.contenido.strip()}\n```"
    if nodo.tipo == "ARBOL":
        return f"```text\n{nodo.contenido.strip()}\n```"
    if nodo.tipo == "TABLA_CRUDA":
        return nodo.contenido
    return nodo.contenido


# ---------------------------------------------------------------------------
# Depuración visual del AST (modo --debug)
# ---------------------------------------------------------------------------
def debug_ast(nodo: Nodo, nivel: int = 0, es_ultimo: bool = True, prefijo: str = "") -> None:
    """Imprime el AST en formato de árbol con caracteres Unicode."""
    if nivel == 0:
        tipo_str = nodo.tipo
        hijos_str = f" ({len(nodo.hijos)} hijos)" if nodo.hijos else ""
        print(f"\n🌳 {tipo_str}{hijos_str}")
        for i, hijo in enumerate(nodo.hijos):
            es_ultimo_hijo = (i == len(nodo.hijos) - 1)
            debug_ast(hijo, nivel + 1, es_ultimo_hijo, "")
        return

    conector = "└── " if es_ultimo else "├── "
    tipo_str = nodo.tipo
    contenido_preview = nodo.contenido[:55].replace("\n", "\\n")
    meta_info = ""
    if nodo.meta:
        meta_info = f" [{', '.join(f'{k}={v}' for k, v in nodo.meta.items())}]"
    print(f"{prefijo}{conector}{tipo_str}: {contenido_preview}{meta_info}")

    if es_ultimo:
        prefijo_hijos = prefijo + "    "
    else:
        prefijo_hijos = prefijo + "│   "

    for i, hijo in enumerate(nodo.hijos):
        es_ultimo_hijo = (i == len(nodo.hijos) - 1)
        debug_ast(hijo, nivel + 1, es_ultimo_hijo, prefijo_hijos)


# ---------------------------------------------------------------------------
# Metadatos (tecnologías, licencia, badges)
# ---------------------------------------------------------------------------
def detectar_tecnologias(texto: str) -> List[Tuple[str, str, str]]:
    """Busca palabras clave de tecnologías en el texto."""
    encontradas = []
    for clave, (nombre, color, logo) in BADGES_TECNOLOGIAS.items():
        if re.search(rf"\b{re.escape(clave)}\b", texto.lower()):
            encontradas.append((nombre, color, logo))
    logging.debug("Tecnologías detectadas: %s", encontradas)
    return encontradas


def detectar_licencia(
    texto: str, manual: Optional[str] = None
) -> Tuple[Optional[str], Optional[str]]:
    """Devuelve la licencia detectada en el texto o la indicada manualmente."""
    if manual:
        for clave, (nombre, color) in LICENCIAS_CONOCIDAS.items():
            if clave == manual.lower() or nombre.lower() == manual.lower():
                return nombre, color
        return manual, "lightgrey"
    for clave, (nombre, color) in LICENCIAS_CONOCIDAS.items():
        if clave in texto.lower():
            logging.debug("Licencia detectada: %s", nombre)
            return nombre, color
    return None, None


def generar_badges(tecnologias, licencia) -> str:
    """Arma la línea de badges de Shields.io."""
    badges = []
    for nombre, color, logo in tecnologias:
        nombre_esc = nombre.replace(" ", "%20")
        badges.append(
            f"![](https://img.shields.io/badge/{nombre_esc}-{color}?style=for-the-badge&logo={logo}&logoColor=white)"
        )
    if licencia[0]:
        nombre_lic, color_lic = licencia
        nombre_lic_esc = nombre_lic.replace(" ", "%20")
        badges.append(
            f"![](https://img.shields.io/badge/License-{nombre_lic_esc}-{color_lic}?style=for-the-badge)"
        )
    return " ".join(badges)


# ---------------------------------------------------------------------------
# Tabla de contenidos
# ---------------------------------------------------------------------------
def normalizar_ancla(texto: str) -> str:
    """Convierte un título en el fragmento de enlace que usa GitHub."""
    limpio = re.sub(r":[a-z_]+:", "", texto)
    limpio = re.sub(r"[^\w\s-]", "", limpio)
    return re.sub(r"-+", "-", re.sub(r"\s+", "-", limpio.strip().lower())).strip("-")


def generar_toc(ast: Nodo) -> str:
    """Genera la TOC a partir de los nodos SECTION del AST."""
    secciones = []

    def recolectar(nodo, dentro_codigo=False):
        if nodo.tipo == "BLOQUE_CODIGO":
            for h in nodo.hijos:
                recolectar(h, True)
            return
        if nodo.tipo == "SECCION" and not dentro_codigo:
            match = re.match(r"^(#+)\s(.*)", nodo.contenido)
            if match:
                titulo = match.group(2).strip()
                if titulo and len(titulo.split()) <= 8 and "|" not in titulo and "©" not in titulo:
                    secciones.append(titulo)
        for h in nodo.hijos:
            recolectar(h, dentro_codigo)

    recolectar(ast)
    if not secciones:
        return ""

    lineas = ["## 📋 Tabla de Contenidos"]
    for titulo in secciones:
        ancla = normalizar_ancla(titulo)
        lineas.append(f"- [{titulo}](#{ancla})")
    return "\n".join(lineas) + "\n"


# ---------------------------------------------------------------------------
# Pipeline principal
# ---------------------------------------------------------------------------
def construir_readme(contenido: str, argumentos) -> str:
    """Orquesta la generación completa del README.md."""
    logging.info("Iniciando generación del README.")

    texto = sanitizar(contenido)

    titulo = "Nombre del Proyecto"
    primera = texto.splitlines()[0].strip() if texto else ""
    if primera.startswith("# "):
        titulo = primera[2:].strip()
        texto = "\n".join(texto.splitlines()[1:]).lstrip("\n")

    texto = forzar_titulos(texto)
    texto = envolver_comandos(texto)

    lineas = texto.splitlines()
    bloques = detectar_bloques(lineas)
    ast = construir_ast(bloques)

    if getattr(argumentos, "debug", False):
        debug_ast(ast)

    walk(ast, visitor_emojis)
    ast = normalizar_ast(ast)

    if not validar_estructura(ast):
        logging.warning("La estructura del AST es débil; el resultado puede ser incompleto.")

    cuerpo = renderizar_md(ast)
    toc = generar_toc(ast) if not getattr(argumentos, "no_toc", False) else ""

    tecnologias = detectar_tecnologias(contenido)
    info_licencia = detectar_licencia(contenido, argumentos.license)
    badges = generar_badges(tecnologias, info_licencia)

    logo_md = ""
    if argumentos.logo:
        logo_md = f'<p align="center">\n  <img src="{argumentos.logo}" alt="Logo" width="200"/>\n</p>\n'
    else:
        for ruta in [Path("Fotos/logo.png"), Path("Fotos/logo.jpg"), Path("logo.png")]:
            if ruta.exists():
                logo_rel = os.path.relpath(str(ruta), os.getcwd())
                logo_md = f'<p align="center">\n  <img src="{logo_rel}" alt="Logo" width="200"/>\n</p>\n'
                break

    partes = [f"# {titulo}", ""]
    if logo_md:
        partes.append(logo_md)
    if badges:
        partes.append(badges + "\n")
    if toc:
        partes.append(toc + "\n")
    partes.append(cuerpo)

    if not getattr(argumentos, "no_creditos", False):
        partes.append("\n## ⭐ Créditos")
        partes.append("Creado con la herramienta de generación de README.")

    readme = "\n".join(partes)
    readme = reparar_fences(readme)
    readme = re.sub(r"\n{3,}", "\n\n", readme)
    logging.info("README generado exitosamente.")
    return readme.rstrip() + "\n"


# ---------------------------------------------------------------------------
# Interfaz de línea de comandos
# ---------------------------------------------------------------------------
def main() -> None:
    configurar_logging()
    try:
        parser = argparse.ArgumentParser(description="Generador de README.md para GitHub")
        sub = parser.add_subparsers(dest="comando")
        build = sub.add_parser("build", help="Generar README.md")
        grupo = build.add_mutually_exclusive_group()
        grupo.add_argument("--txt", help="Ruta al archivo .txt")
        grupo.add_argument("--docx", help="Ruta al archivo .docx")
        grupo.add_argument("--auto", action="store_true", help="Detectar archivo automáticamente")
        build.add_argument("-o", "--output", default="README.md", help="Archivo de salida")
        build.add_argument("--license", help="Licencia del proyecto (MIT, GPL, etc.)")
        build.add_argument("--logo", help="Ruta al archivo de logo")
        build.add_argument("--no-toc", action="store_true", help="Omitir tabla de contenidos")
        build.add_argument("--no-credits", action="store_true", help="Omitir sección de créditos")
        build.add_argument("--no-mandatory", action="store_true", help="No insertar secciones obligatorias faltantes")
        build.add_argument("--debug", action="store_true", help="Mostrar el AST generado en consola para depuración")

        args = parser.parse_args()

        if args.comando == "build":
            if args.auto or (not args.txt and not args.docx):
                ruta = buscar_archivo()
                if not ruta:
                    logging.error("No se encontró archivo fuente.")
                    return
                if ruta.suffix == ".txt":
                    args.txt = str(ruta)
                else:
                    args.docx = str(ruta)

            # Validación de entrada
            if args.txt:
                validado = validar_archivo(args.txt)
                if not validado:
                    return
                args.txt = validado
            if args.docx:
                validado = validar_archivo(args.docx)
                if not validado:
                    return
                args.docx = validado

            contenido = leer_txt(args.txt) if args.txt else leer_docx(args.docx) if args.docx else None
            if not contenido:
                logging.error("No se pudo leer el archivo fuente.")
                return

            readme = construir_readme(contenido, args)
            Path(args.output).write_text(readme, encoding="utf-8")
            print(f"✅ README generado: {args.output}")

    except KeyboardInterrupt:
        print("\n🛑 Operación cancelada por el usuario.")
        sys.exit(0)
    except Exception as error:
        logging.exception("Error inesperado al generar el README.")
        print(f"❌ Error: {error}")
        sys.exit(1)


if __name__ == "__main__":
    main()
